# -*- coding: utf-8 -*-
"""
Erzeugt eine premium Angebotsvorlage (.docx) im Design von webdesign.evglab.com:
minimalistisch, editorial, viel Weißraum, Geist-Font, schwarze Akzent-Blöcke.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- Palette ----
INK        = RGBColor(0x0A, 0x0A, 0x0A)   # primary text / dark blocks
SECONDARY  = RGBColor(0x6B, 0x6B, 0x6B)   # muted gray
FAINT      = RGBColor(0x9A, 0x9A, 0x9A)   # very light labels
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LINE       = "E6E6E6"                       # hairline hex (no #)
SURFACE    = "F5F5F5"                       # subtle surface fill
DARK_HEX   = "0A0A0A"

FONT       = "Geist"        # Marken-Font; Word ersetzt sauber, falls nicht installiert
FONT_FB    = "Segoe UI"     # Fallback wird via Theme nicht erzwungen, nur Referenz

# --------------------------------------------------------------- Helpers -----

def set_run(run, *, size=11, color=INK, bold=False, font=FONT, spacing=None, caps=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), font)
    if spacing is not None:
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(int(spacing)))
        rpr.append(sp)
    if caps:
        c = OxmlElement('w:caps')
        c.set(qn('w:val'), '1')
        rpr.append(c)
    return run


def para(container, text="", *, size=11, color=INK, bold=False, align=None,
         space_before=0, space_after=6, line=1.4, caps=False, spacing=None, font=FONT):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if text:
        r = p.add_run(text)
        set_run(r, size=size, color=color, bold=bold, caps=caps, spacing=spacing, font=font)
    return p


def hairline(paragraph, color=LINE, size=6, space_after=0):
    """Untere Haarlinie an einem Absatz."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    pPr.append(pbdr)
    paragraph.paragraph_format.space_after = Pt(space_after)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right),
                     ('left', left), ('right', right)):
        e = OxmlElement(f'w:{tag}')
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def cell_borders(cell, *, bottom=None, top=None, left=None, right=None):
    """bottom/top/left/right = (hex, size) oder None."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for name, spec in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement(f'w:{name}')
        if spec is None:
            e.set(qn('w:val'), 'nil')
        else:
            hexc, sz = spec
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), str(sz))
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), hexc)
        borders.append(e)
    tcPr.append(borders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'nil')
        borders.append(e)
    tblPr.append(borders)


def cell_text(cell, text, *, size=11, color=INK, bold=False, align=None,
              caps=False, spacing=None, space_after=0, valign=None, font=FONT, line=1.3):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold, caps=caps, spacing=spacing, font=font)
    if valign is not None:
        cell.vertical_alignment = valign
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldStart = OxmlElement('w:fldChar'); fldStart.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldStart); run._r.append(instr); run._r.append(fldEnd)
    set_run(run, size=8, color=FAINT)


def spacer(container, pts=12):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("")
    set_run(r, size=int(pts))
    return p


# --------------------------------------------------------------- Document ----

doc = Document()

# Core properties
doc.core_properties.title = "Angebot – Erik EvgLab"
doc.core_properties.author = "Erik von Gregory"
doc.core_properties.company = "Erik EvgLab"

# Normal style
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.space_after = Pt(6)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
    rfonts.set(qn(attr), FONT)

# Page geometry (A4, großzügige Ränder)
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.4)
section.right_margin = Cm(2.4)

CONTENT_W = section.page_width - section.left_margin - section.right_margin


# ---------------------------------------------------------- Footer (alle) ----

def build_footer(sec):
    footer = sec.footer
    footer.is_linked_to_previous = False
    # vorhandenen leeren Absatz nutzen
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    hairline(p, color=LINE, size=4, space_after=4)
    # Kontaktzeile
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    set_run(p2.add_run("Erik EvgLab  ·  Hauptstraße 18, 86925 Fuchstal  ·  info@evglab.com  ·  +49 173 170 6012  ·  webdesign.evglab.com"),
            size=8, color=SECONDARY)
    p3 = footer.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    add_page_number(p3)

build_footer(section)


# =============================================================== COVER ========

# Kopfzeile: Badge "E" | Wortmarke | ANGEBOT  (eine Tabelle, keine Verschachtelung)
head = doc.add_table(rows=1, cols=3)
head.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(head)
head.autofit = False
head.columns[0].width = Cm(1.0)
head.columns[1].width = Cm(8.0)
head.columns[2].width = Cm(7.2)

badge = head.cell(0, 0)
shade_cell(badge, DARK_HEX)
set_cell_margins(badge, top=60, bottom=60, left=60, right=60)
badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
cell_text(badge, "E", size=13, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          valign=WD_ALIGN_VERTICAL.CENTER)

nm = head.cell(0, 1)
nm.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_margins(nm, top=60, bottom=60, left=160, right=80)
cell_text(nm, "Erik EvgLab", size=13, bold=True, valign=WD_ALIGN_VERTICAL.CENTER)

right = head.cell(0, 2)
right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
cell_text(right, "ANGEBOT", size=12, color=SECONDARY, bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT, caps=True, spacing=60,
          valign=WD_ALIGN_VERTICAL.CENTER)

spacer(doc, 36)

# Eyebrow
para(doc, "WEBDESIGN & ENTWICKLUNG", size=9, color=FAINT, caps=True, spacing=80, space_after=10)

# Großer Titel
t = doc.add_paragraph()
t.paragraph_format.space_after = Pt(4)
t.paragraph_format.line_spacing = 1.02
set_run(t.add_run("Dein neuer\nWebauftritt."), size=40, bold=True, color=INK)

para(doc, "Ein Angebot, das so durchdacht ist wie die Website, die du dafür bekommst.",
     size=13, color=SECONDARY, space_before=2, space_after=0, line=1.45)

spacer(doc, 30)

# Meta-Karte (Angebot-Nr, Datum, Gültig bis, Kunde)
meta = doc.add_table(rows=2, cols=4)
remove_table_borders(meta)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ["ANGEBOTS-NR.", "DATUM", "GÜLTIG BIS", "PROJEKT"]
values = ["2026-0001", "TT.MM.JJJJ", "TT.MM.JJJJ", "[Projektname]"]
for i in range(4):
    lc = meta.cell(0, i)
    cell_borders(lc, top=(LINE, 8))
    set_cell_margins(lc, top=120, bottom=20, left=0, right=80)
    cell_text(lc, labels[i], size=8, color=FAINT, caps=True, spacing=60)
    vc = meta.cell(1, i)
    set_cell_margins(vc, top=0, bottom=120, left=0, right=80)
    cell_text(vc, values[i], size=12, bold=True)

spacer(doc, 26)

# Kunde / Anbieter Block
party = doc.add_table(rows=1, cols=2)
remove_table_borders(party)
party.alignment = WD_TABLE_ALIGNMENT.CENTER
party.columns[0].width = CONTENT_W // 2
party.columns[1].width = CONTENT_W // 2

c_to = party.cell(0, 0)
set_cell_margins(c_to, top=140, bottom=140, left=140, right=140)
shade_cell(c_to, SURFACE)
cell_text(c_to, "ANGEBOT FÜR", size=8, color=FAINT, caps=True, spacing=60, space_after=6)
for line, b in [("[Firmenname]", True), ("[Ansprechpartner]", False),
                ("[Straße & Hausnummer]", False), ("[PLZ & Ort]", False),
                ("[E-Mail]", False)]:
    p = c_to.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run(line), size=11, bold=b, color=INK if b else SECONDARY)

c_from = party.cell(0, 1)
set_cell_margins(c_from, top=140, bottom=140, left=140, right=140)
cell_borders(c_from, top=(LINE, 8), bottom=(LINE, 8), left=(LINE, 8), right=(LINE, 8))
cell_text(c_from, "ANBIETER", size=8, color=FAINT, caps=True, spacing=60, space_after=6)
for line, b in [("Erik von Gregory · EvgLab", True), ("Hauptstraße 18", False),
                ("86925 Fuchstal", False), ("info@evglab.com", False),
                ("+49 173 170 6012", False)]:
    p = c_from.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run(line), size=11, bold=b, color=INK if b else SECONDARY)

doc.add_page_break()


# =========================================================== ANSCHREIBEN ======

def section_title(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(num + "  "); set_run(r1, size=11, color=FAINT, bold=True)
    r2 = p.add_run(title); set_run(r2, size=16, color=INK, bold=True)
    hp = doc.add_paragraph(); hp.paragraph_format.space_after = Pt(14)
    hairline(hp, color=LINE, size=6, space_after=0)

section_title("01", "Worum es geht")

para(doc, "Hallo [Vorname],", size=12, bold=True, space_after=8)
para(doc,
     "danke für dein Interesse und das offene Gespräch. Auf den nächsten Seiten findest du ein "
     "konkretes Angebot für deinen neuen Webauftritt – klar aufgeschlüsselt, ohne Kleingedrucktes "
     "und ohne Vorlagen-Look von der Stange.",
     size=11, color=SECONDARY, space_after=8)
para(doc,
     "Du bekommst keine Website, die aussieht wie tausend andere. Du bekommst einen Auftritt mit "
     "Charakter, der zu dir passt, schnell lädt und am Ende eines tut: aus Besuchern Anfragen machen. "
     "Alles aus einer Hand – Konzept, Design und Entwicklung.",
     size=11, color=SECONDARY, space_after=18)

# Highlight-Zeile (3 Mini-Kennzahlen)
stats = doc.add_table(rows=1, cols=3)
remove_table_borders(stats)
stats.alignment = WD_TABLE_ALIGNMENT.CENTER
facts = [("7 Tage", "bis zur fertigen Seite"), ("100 %", "individuell, kein Template"),
         ("1 : 1", "direkter Ansprechpartner")]
for i, (big, small) in enumerate(facts):
    c = stats.cell(0, i)
    set_cell_margins(c, top=140, bottom=140, left=140, right=140)
    shade_cell(c, SURFACE)
    cell_text(c, big, size=20, bold=True, space_after=2)
    p = c.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(small), size=9.5, color=SECONDARY)
    if i < 2:
        # kleine Lücke zwischen den Zellen über Zellmargin rechts simulieren
        pass

spacer(doc, 20)


# =========================================================== LEISTUNGEN ========

section_title("02", "Leistungsumfang")

para(doc,
     "Die folgenden Positionen bilden den Kern des Projekts. Du wählst, was du brauchst – "
     "alles ist einzeln kalkuliert.",
     size=11, color=SECONDARY, space_after=14)

items = [
    ("Komplette Website",
     "Mehrseitiger Auftritt: Konzept, individuelles Design, Entwicklung (Next.js), "
     "responsiv für alle Geräte, Grund-SEO & Kontaktformular.",
     "2.500 €"),
    ("Landingpage",
     "Eine fokussierte Seite mit klarer Conversion-Strategie – ideal für Kampagnen, "
     "ein Produkt oder einen einzelnen Service.",
     "1.500 €"),
    ("Laufende Betreuung",
     "Updates, kleine Anpassungen, Hosting-Monitoring & Erreichbarkeit. "
     "Monatlich kündbar.",
     "99 € / Monat"),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
tbl.columns[0].width = Cm(4.6)
tbl.columns[1].width = Cm(8.7)
tbl.columns[2].width = Cm(2.9)

# Kopfzeile (dunkel)
hdr = tbl.rows[0].cells
heads = ["Leistung", "Beschreibung", "Investition"]
aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]
for i, c in enumerate(hdr):
    shade_cell(c, DARK_HEX)
    set_cell_margins(c, top=100, bottom=100, left=120, right=120)
    cell_text(c, heads[i], size=9, color=WHITE, bold=True, caps=True, spacing=40, align=aligns[i],
              valign=WD_ALIGN_VERTICAL.CENTER)

for name, desc, price in items:
    row = tbl.add_row().cells
    for i, c in enumerate(row):
        set_cell_margins(c, top=120, bottom=120, left=120, right=120)
        cell_borders(c, bottom=(LINE, 6))
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_text(row[0], name, size=11, bold=True)
    cell_text(row[1], desc, size=10, color=SECONDARY, line=1.35)
    cell_text(row[2], price, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

spacer(doc, 16)

# Summenblock (rechtsbündig)
sumt = doc.add_table(rows=3, cols=2)
remove_table_borders(sumt)
sumt.alignment = WD_TABLE_ALIGNMENT.RIGHT
sumt.columns[0].width = Cm(9)
sumt.columns[1].width = Cm(4)

def sum_row(idx, label, value, *, strong=False, dark=False):
    lc = sumt.cell(idx, 0); vc = sumt.cell(idx, 1)
    for c in (lc, vc):
        set_cell_margins(c, top=90, bottom=90, left=120, right=120)
        if dark:
            shade_cell(c, DARK_HEX)
    col = WHITE if dark else (INK if strong else SECONDARY)
    cell_text(lc, label, size=11 if not strong else 12, color=col, bold=strong,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
    cell_text(vc, value, size=11 if not strong else 13, color=col, bold=True,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

sum_row(0, "Zwischensumme", "[Betrag] €")
sum_row(1, "Umsatzsteuer (§ 19 UStG, Kleinunternehmer)", "0,00 €")
sum_row(2, "Gesamtbetrag", "[Betrag] €", strong=True, dark=True)

para(doc, "Es wird gemäß § 19 UStG keine Umsatzsteuer ausgewiesen.",
     size=9, color=FAINT, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=6, space_after=0)

doc.add_page_break()


# =============================================================== ABLAUF ========

section_title("03", "So läuft die Zusammenarbeit")

steps = [
    ("Kennenlernen", "Wir klären Ziele, Zielgruppe und Anforderungen – kostenlos und unverbindlich."),
    ("Konzept & Struktur", "Ich entwickle Aufbau und Inhaltsgerüst, abgestimmt auf deine Botschaft."),
    ("Design", "Du bekommst ein individuelles Design – kein Baukasten, sondern dein Auftritt."),
    ("Entwicklung", "Saubere, schnelle Umsetzung mit moderner Technik und Fokus auf Performance."),
    ("Launch & Betreuung", "Wir gehen live. Auf Wunsch bleibe ich danach an deiner Seite."),
]
for i, (title, desc) in enumerate(steps, 1):
    row = doc.add_table(rows=1, cols=2)
    remove_table_borders(row)
    row.columns[0].width = Cm(1.4)
    row.columns[1].width = CONTENT_W - Cm(1.4)
    nc = row.cell(0, 0)
    nc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_text(nc, f"{i:02d}", size=15, color=FAINT, bold=True)
    tc = row.cell(0, 1)
    cell_text(tc, title, size=12, bold=True, space_after=2)
    p = tc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.35
    set_run(p.add_run(desc), size=10.5, color=SECONDARY)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(0)
    hairline(sp, color=LINE, size=4, space_after=0)
    spacer(doc, 6)

spacer(doc, 12)


# =========================================================== KONDITIONEN =======

section_title("04", "Konditionen")

conds = [
    ("Lieferzeit", "In der Regel 7 Tage ab Bereitstellung aller Inhalte (Texte, Bilder, Logo)."),
    ("Zahlung", "50 % bei Auftragsbestätigung, 50 % nach Freigabe vor dem Launch."),
    ("Gültigkeit", "Dieses Angebot ist 14 Tage ab Angebotsdatum gültig."),
    ("Leistungen Dritter", "Hosting, Domain und externe Lizenzen werden – falls nötig – transparent weitergegeben."),
    ("Nutzungsrechte", "Mit vollständiger Bezahlung erhältst du umfassende Nutzungsrechte an den für dich "
     "erstellten Inhalten – du darfst die Website uneingeschränkt nutzen, veröffentlichen und bearbeiten. "
     "Das Urheberrecht verbleibt gemäß deutschem Recht beim Ersteller; eingesetzte Fremd-Komponenten "
     "(Open-Source, Schriften, Stockbilder) behalten ihre jeweiligen Lizenzen."),
]
for label, text in conds:
    row = doc.add_table(rows=1, cols=2)
    remove_table_borders(row)
    row.columns[0].width = Cm(4.2)
    row.columns[1].width = CONTENT_W - Cm(4.2)
    lc = row.cell(0, 0)
    set_cell_margins(lc, top=60, bottom=60, left=0, right=120)
    cell_text(lc, label, size=10.5, bold=True)
    rc = row.cell(0, 1)
    set_cell_margins(rc, top=60, bottom=60, left=0, right=0)
    cell_text(rc, text, size=10.5, color=SECONDARY, line=1.35)

doc.add_page_break()


# ============================================================= ANNAHME =========

section_title("05", "Angebot annehmen")

para(doc,
     "Du möchtest loslegen? Dann unterschreib einfach unten und schick mir das Dokument zurück – "
     "oder antworte mit einem kurzen „Passt, los geht’s“ per E-Mail. Ich melde mich umgehend mit "
     "den nächsten Schritten.",
     size=11, color=SECONDARY, space_after=20)

# Auswahl-Kästchen
choice = doc.add_table(rows=1, cols=1)
remove_table_borders(choice)
opts = ["Komplette Website – 2.500 €",
        "Landingpage – 1.500 €",
        "+ Laufende Betreuung – 99 € / Monat"]
cc = choice.cell(0, 0)
shade_cell(cc, SURFACE)
set_cell_margins(cc, top=140, bottom=140, left=140, right=140)
cell_text(cc, "GEWÜNSCHTE LEISTUNG", size=8, color=FAINT, caps=True, spacing=60, space_after=8)
for o in opts:
    p = cc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run("☐  "), size=12, color=INK)
    set_run(p.add_run(o), size=11, color=INK)

spacer(doc, 30)

# Unterschriftsfelder
sign = doc.add_table(rows=2, cols=2)
remove_table_borders(sign)
sign.columns[0].width = CONTENT_W // 2
sign.columns[1].width = CONTENT_W // 2

for col, label in [(0, "Auftraggeber:in"), (1, "Erik von Gregory · EvgLab")]:
    line_cell = sign.cell(0, col)
    set_cell_margins(line_cell, top=200, bottom=20, left=0, right=200)
    cell_borders(line_cell, bottom=(DARK_HEX, 8))
    line_cell.text = ""
    lab_cell = sign.cell(1, col)
    set_cell_margins(lab_cell, top=20, bottom=0, left=0, right=200)
    cell_text(lab_cell, "Ort, Datum, Unterschrift", size=8.5, color=FAINT, space_after=2)
    p = lab_cell.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label), size=10, color=SECONDARY)

spacer(doc, 36)

# Abschluss-Banner (dunkel)
banner = doc.add_table(rows=1, cols=1)
remove_table_borders(banner)
bc = banner.cell(0, 0)
shade_cell(bc, DARK_HEX)
set_cell_margins(bc, top=200, bottom=200, left=200, right=200)
cell_text(bc, "Bereit, wenn du es bist.", size=18, color=WHITE, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
p = bc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("info@evglab.com   ·   +49 173 170 6012   ·   webdesign.evglab.com"),
        size=10.5, color=RGBColor(0xCF, 0xCF, 0xCF))


# ---------------------------------------------------------------- Save --------
out = "Angebotsvorlage_EvgLab.docx"
doc.save(out)
print("OK ->", out)
