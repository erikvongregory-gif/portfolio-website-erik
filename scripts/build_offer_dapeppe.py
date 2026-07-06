# -*- coding: utf-8 -*-
"""
Konkretes Angebot fuer "Da Peppe" mit zwei Optionen:
  A) Komplette Website  – 2.500 € einmalig
  B) Komplett + Betreuung – 1.500 € einmalig + 99 €/Monat
Design im Look von webdesign.evglab.com.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- Palette ----
INK        = RGBColor(0x0A, 0x0A, 0x0A)
SECONDARY  = RGBColor(0x6B, 0x6B, 0x6B)
FAINT      = RGBColor(0x9A, 0x9A, 0x9A)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LINE       = "E6E6E6"
SURFACE    = "F5F5F5"
DARK_HEX   = "0A0A0A"
FONT       = "Geist"  # als Desktop-Schrift installiert (scripts/install_geist_font.ps1)

# Angebotsdaten
DATUM        = "25.06.2026"
GUELTIG_BIS  = "09.07.2026"
ANGEBOT_NR   = "DP-2026-001"

# --------------------------------------------------------------- Helpers -----

def set_run(run, *, size=11, color=INK, bold=False, font=FONT, spacing=None, caps=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), font)
    if spacing is not None:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(int(spacing))); rpr.append(sp)
    if caps:
        c = OxmlElement('w:caps'); c.set(qn('w:val'), '1'); rpr.append(c)
    return run


def para(container, text="", *, size=11, color=INK, bold=False, align=None,
         space_before=0, space_after=6, line=1.4, caps=False, spacing=None, font=FONT):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after); pf.line_spacing = line
    if text:
        set_run(p.add_run(text), size=size, color=color, bold=bold, caps=caps, spacing=spacing, font=font)
    return p


def hairline(paragraph, color=LINE, size=6, space_after=0):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '6'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom); pPr.append(pbdr)
    paragraph.paragraph_format.space_after = Pt(space_after)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right),
                     ('left', left), ('right', right)):
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)


def cell_borders(cell, *, bottom=None, top=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for name, spec in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement(f'w:{name}')
        if spec is None:
            e.set(qn('w:val'), 'nil')
        else:
            hexc, sz = spec
            e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
            e.set(qn('w:space'), '0'); e.set(qn('w:color'), hexc)
        borders.append(e)
    tcPr.append(borders)


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'nil'); borders.append(e)
    tblPr.append(borders)


def cell_text(cell, text, *, size=11, color=INK, bold=False, align=None,
              caps=False, spacing=None, space_after=0, valign=None, font=FONT, line=1.3):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = line
    set_run(p.add_run(text), size=size, color=color, bold=bold, caps=caps, spacing=spacing, font=font)
    if valign is not None:
        cell.vertical_alignment = valign
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    b = OxmlElement('w:instrText'); b.set(qn('xml:space'), 'preserve'); b.text = 'PAGE'
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'end')
    run._r.append(a); run._r.append(b); run._r.append(c)
    set_run(run, size=8, color=FAINT)


def spacer(container, pts=12):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    set_run(p.add_run(""), size=int(pts))
    return p


def bullet(cell, text, *, bold_word=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run("— "), size=10.5, color=INK, bold=True)
    set_run(p.add_run(text), size=10.5, color=SECONDARY)
    return p


# --------------------------------------------------------------- Document ----
doc = Document()
doc.core_properties.title = "Angebot Website – Da Peppe"
doc.core_properties.author = "Erik von Gregory"
doc.core_properties.company = "Erik EvgLab"

normal = doc.styles['Normal']
normal.font.name = FONT; normal.font.size = Pt(11); normal.font.color.rgb = INK
normal.paragraph_format.line_spacing = 1.4; normal.paragraph_format.space_after = Pt(6)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
    rfonts.set(qn(attr), FONT)

section = doc.sections[0]
section.page_height = Cm(29.7); section.page_width = Cm(21.0)
section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.4); section.right_margin = Cm(2.4)
CONTENT_W = section.page_width - section.left_margin - section.right_margin

# Footer
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_before = Pt(6)
hairline(fp, color=LINE, size=4, space_after=4)
fp2 = footer.add_paragraph(); fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp2.paragraph_format.space_after = Pt(2)
set_run(fp2.add_run("Erik EvgLab  ·  Hauptstraße 18, 86925 Fuchstal  ·  info@evglab.com  ·  +49 173 170 6012  ·  webdesign.evglab.com"),
        size=8, color=SECONDARY)
fp3 = footer.add_paragraph(); fp3.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp3.paragraph_format.space_after = Pt(0)
add_page_number(fp3)


# =============================================================== COVER ========
head = doc.add_table(rows=1, cols=3)
head.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(head); head.autofit = False
head.columns[0].width = Cm(1.0); head.columns[1].width = Cm(8.0); head.columns[2].width = Cm(7.2)
badge = head.cell(0, 0); shade_cell(badge, DARK_HEX)
set_cell_margins(badge, top=60, bottom=60, left=60, right=60)
cell_text(badge, "E", size=13, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, valign=WD_ALIGN_VERTICAL.CENTER)
nm = head.cell(0, 1); set_cell_margins(nm, top=60, bottom=60, left=160, right=80)
cell_text(nm, "Erik EvgLab", size=13, bold=True, valign=WD_ALIGN_VERTICAL.CENTER)
rt = head.cell(0, 2)
cell_text(rt, "ANGEBOT", size=12, color=SECONDARY, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
          caps=True, spacing=60, valign=WD_ALIGN_VERTICAL.CENTER)

spacer(doc, 34)
para(doc, "WEBDESIGN FÜR GASTRONOMIE", size=9, color=FAINT, caps=True, spacing=80, space_after=10)

t = doc.add_paragraph(); t.paragraph_format.space_after = Pt(4); t.paragraph_format.line_spacing = 1.02
set_run(t.add_run("Die Website für\nDa Peppe."), size=40, bold=True, color=INK)
para(doc, "Damit Gäste euch online genauso überzeugend finden, wie euer Essen schmeckt.",
     size=13, color=SECONDARY, space_before=2, space_after=0, line=1.45)

spacer(doc, 28)
meta = doc.add_table(rows=2, cols=4); remove_table_borders(meta); meta.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ["ANGEBOTS-NR.", "DATUM", "GÜLTIG BIS", "PROJEKT"]
values = [ANGEBOT_NR, DATUM, GUELTIG_BIS, "Website Da Peppe"]
for i in range(4):
    lc = meta.cell(0, i); cell_borders(lc, top=(LINE, 8)); set_cell_margins(lc, top=120, bottom=20, left=0, right=80)
    cell_text(lc, labels[i], size=8, color=FAINT, caps=True, spacing=60)
    vc = meta.cell(1, i); set_cell_margins(vc, top=0, bottom=120, left=0, right=80)
    cell_text(vc, values[i], size=12, bold=True)

spacer(doc, 24)
party = doc.add_table(rows=1, cols=2); remove_table_borders(party); party.alignment = WD_TABLE_ALIGNMENT.CENTER
party.columns[0].width = CONTENT_W // 2; party.columns[1].width = CONTENT_W // 2
c_to = party.cell(0, 0); set_cell_margins(c_to, top=140, bottom=140, left=140, right=140); shade_cell(c_to, SURFACE)
cell_text(c_to, "ANGEBOT FÜR", size=8, color=FAINT, caps=True, spacing=60, space_after=6)
for line, b in [("Da Peppe", True), ("z. Hd. Inhaber", False),
                ("[Straße & Hausnummer]", False), ("[PLZ & Ort]", False)]:
    p = c_to.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run(line), size=11, bold=b, color=INK if b else SECONDARY)
c_from = party.cell(0, 1); set_cell_margins(c_from, top=140, bottom=140, left=140, right=140)
cell_borders(c_from, top=(LINE, 8), bottom=(LINE, 8), left=(LINE, 8), right=(LINE, 8))
cell_text(c_from, "ANBIETER", size=8, color=FAINT, caps=True, spacing=60, space_after=6)
for line, b in [("Erik von Gregory · EvgLab", True), ("Hauptstraße 18", False),
                ("86925 Fuchstal", False), ("info@evglab.com", False), ("+49 173 170 6012", False)]:
    p = c_from.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run(line), size=11, bold=b, color=INK if b else SECONDARY)

doc.add_page_break()


# ----------------------------------------------------- section title helper ---
def section_title(num, title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(num + "  "), size=11, color=FAINT, bold=True)
    set_run(p.add_run(title), size=16, color=INK, bold=True)
    hp = doc.add_paragraph(); hp.paragraph_format.space_after = Pt(14); hairline(hp, color=LINE, size=6, space_after=0)


# =========================================================== ANSCHREIBEN ======
section_title("01", "Hallo Peppe")
para(doc,
     "danke für das nette Gespräch und das gute Essen! Wie versprochen, hier mein Angebot für "
     "eure neue Website. Ziel ist einfach: Wer hungrig ist und nach einem guten Italiener sucht, "
     "soll bei euch landen – und sofort Lust bekommen, vorbeizukommen oder zu bestellen.",
     size=11, color=SECONDARY, space_after=8)
para(doc,
     "Ihr bekommt keine Baukasten-Seite von der Stange, sondern einen Auftritt mit Charakter, der "
     "zu Da Peppe passt: schnell, auf dem Handy perfekt und bei Google gut auffindbar. "
     "Alles aus einer Hand – und in der Regel in nur 7 Tagen fertig.",
     size=11, color=SECONDARY, space_after=18)

stats = doc.add_table(rows=1, cols=3); remove_table_borders(stats); stats.alignment = WD_TABLE_ALIGNMENT.CENTER
facts = [("7 Tage", "bis zur fertigen Seite"),
         ("Mobil first", "die meisten Gäste suchen am Handy"),
         ("Lokal", "bei Google in eurer Region gefunden")]
for i, (big, small) in enumerate(facts):
    c = stats.cell(0, i); set_cell_margins(c, top=140, bottom=140, left=140, right=140); shade_cell(c, SURFACE)
    cell_text(c, big, size=18, bold=True, space_after=2)
    p = c.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(small), size=9.5, color=SECONDARY)

spacer(doc, 22)


# ============================================================ ZWEI OPTIONEN ===
section_title("02", "Deine zwei Möglichkeiten")
para(doc, "Du hast die Wahl – je nachdem, ob du lieber einmal zahlst oder rundum betreut sein möchtest:",
     size=11, color=SECONDARY, space_after=14)

cards = doc.add_table(rows=1, cols=3); remove_table_borders(cards); cards.alignment = WD_TABLE_ALIGNMENT.CENTER
cards.autofit = False
cards.columns[0].width = Cm(7.4); cards.columns[1].width = Cm(0.6); cards.columns[2].width = Cm(7.4)

# ---- Option A ----
a = cards.cell(0, 0)
set_cell_margins(a, top=180, bottom=180, left=180, right=180)
cell_borders(a, top=(LINE, 8), bottom=(LINE, 8), left=(LINE, 8), right=(LINE, 8))
cell_text(a, "OPTION A", size=8, color=FAINT, caps=True, spacing=60, space_after=4)
p = a.add_paragraph(); p.paragraph_format.space_after = Pt(8)
set_run(p.add_run("Komplett"), size=15, bold=True, color=INK)
p = a.add_paragraph(); p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("2.500 €"), size=26, bold=True, color=INK)
p = a.add_paragraph(); p.paragraph_format.space_after = Pt(2)
set_run(p.add_run("einmalig"), size=11, color=SECONDARY)
p = a.add_paragraph(); p.paragraph_format.space_after = Pt(10)
set_run(p.add_run("Keine laufenden Kosten bei mir."), size=10, color=SECONDARY)
hp = a.add_paragraph(); hp.paragraph_format.space_after = Pt(10); hairline(hp, color=LINE, size=6, space_after=0)
for line in ["Komplette, mehrseitige Website", "Individuelles Design – kein Template",
             "Einmal zahlen, dann gehört sie dir", "Inhalte pflegst du selbst"]:
    bullet(a, line)
p = a.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("Ideal, wenn du alles selbst in der Hand haben willst."), size=9.5, color=FAINT)

# ---- Spacer ----
cards.cell(0, 1)  # leere Zelle mit Standardabsatz – nur Abstand

# ---- Option B (empfohlen) ----
b = cards.cell(0, 2)
set_cell_margins(b, top=180, bottom=180, left=180, right=180)
shade_cell(b, SURFACE)
cell_borders(b, top=(DARK_HEX, 14), bottom=(DARK_HEX, 14), left=(DARK_HEX, 14), right=(DARK_HEX, 14))
cell_text(b, "OPTION B · EMPFOHLEN", size=8, color=INK, caps=True, spacing=50, bold=True, space_after=4)
p = b.add_paragraph(); p.paragraph_format.space_after = Pt(8)
set_run(p.add_run("Komplett + Betreuung"), size=15, bold=True, color=INK)
p = b.add_paragraph(); p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("1.500 €"), size=26, bold=True, color=INK)
p = b.add_paragraph(); p.paragraph_format.space_after = Pt(2)
set_run(p.add_run("einmalig  +  99 € / Monat"), size=12, bold=True, color=INK)
p = b.add_paragraph(); p.paragraph_format.space_after = Pt(10)
set_run(p.add_run("Niedriger Start – ich kümmere mich laufend."), size=10, color=SECONDARY)
hp = b.add_paragraph(); hp.paragraph_format.space_after = Pt(10); hairline(hp, color="D8D8D8", size=6, space_after=0)
for line in ["Alles aus Option A", "Laufende Updates & Änderungen",
             "Speisekarte & Aktionen immer aktuell", "Technik, Monitoring & Erreichbarkeit",
             "1.000 € weniger zum Start", "Monatlich kündbar"]:
    bullet(b, line)
p = b.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("Ideal, wenn du dich um nichts kümmern willst."), size=9.5, color=FAINT)

spacer(doc, 10)
para(doc,
     "Hinweis: Bei Option B zahlst du zum Start 1.000 € weniger. Dafür halte ich deine Seite "
     "dauerhaft aktuell und bin dein fester Ansprechpartner – die Betreuung ist monatlich kündbar.",
     size=9.5, color=FAINT, space_after=0, line=1.4)

doc.add_page_break()


# ====================================================== IN BEIDEN ENTHALTEN ===
section_title("03", "Das ist in beiden Optionen drin")
incl = [
    "Individuelles Design, abgestimmt auf Da Peppe – kein Baukasten",
    "Digitale Speisekarte, klar und appetitlich präsentiert",
    "Öffnungszeiten, Anfahrt & Kontakt auf einen Blick",
    "Reservierungs- bzw. Kontaktanfrage direkt über die Seite",
    "Perfekt auf dem Handy – dort suchen die meisten Gäste",
    "Google-Grundoptimierung, damit ihr lokal gefunden werdet",
    "Galerie für Gerichte & Ambiente",
    "Schnelle Ladezeit und moderne, saubere Technik",
]
tbl = doc.add_table(rows=0, cols=2); remove_table_borders(tbl); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.columns[0].width = CONTENT_W // 2; tbl.columns[1].width = CONTENT_W // 2
for i in range(0, len(incl), 2):
    row = tbl.add_row().cells
    for j in range(2):
        c = row[j]; set_cell_margins(c, top=60, bottom=60, left=0, right=160)
        if i + j < len(incl):
            cell_text(c, "", size=10.5)
            pp = c.paragraphs[0]
            set_run(pp.add_run("— "), size=10.5, bold=True, color=INK)
            set_run(pp.add_run(incl[i + j]), size=10.5, color=SECONDARY)
            pp.paragraph_format.line_spacing = 1.3

spacer(doc, 14)


# =============================================================== ABLAUF ========
section_title("04", "So läuft es ab")
steps = [
    ("Kennenlernen", "Wir klären, was Da Peppe ausmacht und wen ihr erreichen wollt."),
    ("Inhalte sammeln", "Du gibst mir Speisekarte, ein paar Fotos und Öffnungszeiten."),
    ("Design & Umsetzung", "Ich gestalte und baue deine Seite – individuell und mobil-stark."),
    ("Feedback & Feinschliff", "Du schaust drüber, wir passen an, bis es sitzt."),
    ("Launch", "Wir gehen live – und deine Gäste finden dich online."),
]
for i, (title, desc) in enumerate(steps, 1):
    row = doc.add_table(rows=1, cols=2); remove_table_borders(row)
    row.columns[0].width = Cm(1.4); row.columns[1].width = CONTENT_W - Cm(1.4)
    cell_text(row.cell(0, 0), f"{i:02d}", size=15, color=FAINT, bold=True)
    tc = row.cell(0, 1); cell_text(tc, title, size=12, bold=True, space_after=2)
    p = tc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.35
    set_run(p.add_run(desc), size=10.5, color=SECONDARY)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(0); hairline(sp, color=LINE, size=4, space_after=0)
    spacer(doc, 6)

spacer(doc, 10)


# =========================================================== KONDITIONEN =======
section_title("05", "Konditionen")
conds = [
    ("Lieferzeit", "In der Regel 7 Tage, sobald alle Inhalte (Speisekarte, Fotos, Texte) vorliegen."),
    ("Zahlung", "50 % bei Auftrag, 50 % nach deiner Freigabe vor dem Launch."),
    ("Betreuung (Option B)", "99 € / Monat, monatlich kündbar. Beginnt mit dem Launch."),
    ("Gültigkeit", "Dieses Angebot ist bis zum " + GUELTIG_BIS + " gültig."),
    ("Nutzungsrechte", "Mit vollständiger Bezahlung erhältst du umfassende Nutzungsrechte an den für dich "
     "erstellten Inhalten – du darfst die Website uneingeschränkt nutzen, veröffentlichen und bearbeiten. "
     "Das Urheberrecht verbleibt gemäß deutschem Recht beim Ersteller; eingesetzte Fremd-Komponenten "
     "(Open-Source, Schriften, Stockbilder) behalten ihre jeweiligen Lizenzen."),
]
for label, text in conds:
    row = doc.add_table(rows=1, cols=2); remove_table_borders(row)
    row.columns[0].width = Cm(4.6); row.columns[1].width = CONTENT_W - Cm(4.6)
    lc = row.cell(0, 0); set_cell_margins(lc, top=60, bottom=60, left=0, right=120)
    cell_text(lc, label, size=10.5, bold=True)
    rc = row.cell(0, 1); set_cell_margins(rc, top=60, bottom=60, left=0, right=0)
    cell_text(rc, text, size=10.5, color=SECONDARY, line=1.35)

doc.add_page_break()


# ============================================================= ANNAHME =========
section_title("06", "Los geht’s")
para(doc,
     "Du bist überzeugt? Dann kreuze einfach deine Option an, unterschreib unten – und ich starte. "
     "Du kannst mir auch einfach kurz Bescheid geben, dann mache ich den Rest.",
     size=11, color=SECONDARY, space_after=20)

choice = doc.add_table(rows=1, cols=1); remove_table_borders(choice)
cc = choice.cell(0, 0); shade_cell(cc, SURFACE); set_cell_margins(cc, top=160, bottom=160, left=160, right=160)
cell_text(cc, "ICH ENTSCHEIDE MICH FÜR", size=8, color=FAINT, caps=True, spacing=60, space_after=10)
for o in ["Option A – Komplett · 2.500 € einmalig",
          "Option B – Komplett + Betreuung · 1.500 € einmalig + 99 € / Monat"]:
    p = cc.add_paragraph(); p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.3
    set_run(p.add_run("☐  "), size=13, color=INK)
    set_run(p.add_run(o), size=11.5, color=INK, bold=True)

spacer(doc, 28)
sign = doc.add_table(rows=2, cols=2); remove_table_borders(sign)
sign.columns[0].width = CONTENT_W // 2; sign.columns[1].width = CONTENT_W // 2
for col, label in [(0, "Da Peppe"), (1, "Erik von Gregory · EvgLab")]:
    lcell = sign.cell(0, col); set_cell_margins(lcell, top=200, bottom=20, left=0, right=200)
    cell_borders(lcell, bottom=(DARK_HEX, 8)); lcell.text = ""
    labcell = sign.cell(1, col); set_cell_margins(labcell, top=20, bottom=0, left=0, right=200)
    cell_text(labcell, "Ort, Datum, Unterschrift", size=8.5, color=FAINT, space_after=2)
    p = labcell.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label), size=10, color=SECONDARY)

spacer(doc, 34)
banner = doc.add_table(rows=1, cols=1); remove_table_borders(banner)
bc = banner.cell(0, 0); shade_cell(bc, DARK_HEX); set_cell_margins(bc, top=200, bottom=200, left=200, right=200)
cell_text(bc, "Freuen uns auf Da Peppe – online.", size=18, color=WHITE, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
p = bc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
set_run(p.add_run("info@evglab.com   ·   +49 173 170 6012   ·   webdesign.evglab.com"),
        size=10.5, color=RGBColor(0xCF, 0xCF, 0xCF))

out = "Angebot_DaPeppe.docx"
try:
    doc.save(out)
except PermissionError:
    import time
    out = f"Angebot_DaPeppe_{time.strftime('%H%M')}.docx"
    doc.save(out)
    print("Hinweis: Originaldatei war gesperrt (in Word offen).")
print("OK ->", out)
